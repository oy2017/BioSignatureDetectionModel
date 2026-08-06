#!/usr/bin/env python3
"""Render the revised manuscript markdown to .docx and .pdf.

One parse, two renderers, so the two outputs cannot drift apart. Handles the
constructs this manuscript actually uses: ATX headings, paragraphs with inline
bold/italic, pipe tables, embedded figures, and the numbered reference list.
"""
import html
import os
import re
import sys
import urllib.parse

SRC = "/mnt/c/Users/owenh/BioSignatureDetectionModel/revision/revised_manuscript.md"
OUTDIR = "/mnt/c/Users/owenh/BioSignatureDetectionModel/revision"
FIGROOT = OUTDIR
DEJAVU = "/usr/share/fonts/truetype/dejavu"


# ---------------------------------------------------------------- parsing

def parse(md):
    """Markdown -> list of blocks: (kind, payload)."""
    blocks, lines, i = [], md.split("\n"), 0
    while i < len(lines):
        ln = lines[i]
        s = ln.strip()

        if not s or s == "---":
            i += 1
            continue

        if s.startswith(">"):
            blocks.append(("quote", re.sub(r"^>\s?", "", s)))
            i += 1
            continue

        m = re.match(r"^(#{1,6})\s+(.*)$", s)
        if m:
            blocks.append(("h", (len(m.group(1)), m.group(2).strip())))
            i += 1
            continue

        m = re.match(r"^!\[[^\]]*\]\(([^)]+)\)\s*$", s)
        if m:
            blocks.append(("img", urllib.parse.unquote(m.group(1))))
            i += 1
            continue

        if s.startswith("|"):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if not all(re.fullmatch(r":?-{2,}:?", c) for c in cells):
                    rows.append(cells)
                i += 1
            if rows:
                blocks.append(("table", rows))
            continue

        m = re.match(r"^(\d+)\.\s+(.*)$", s)
        if m:
            blocks.append(("li", (m.group(1) + ".", m.group(2))))
            i += 1
            continue
        if re.match(r"^[-*]\s+", s):
            blocks.append(("li", ("•", re.sub(r"^[-*]\s+", "", s))))
            i += 1
            continue

        blocks.append(("p", s))
        i += 1
    return blocks


TOKEN = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`|<sup>.*?</sup>|<sub>.*?</sub>)", re.S)


def runs(text):
    """Inline markdown -> [(text, bold, italic, script)].

    `script` is "super", "sub" or None. Superscripts carry the author/affiliation
    markers, so they must render as real superscript rather than being stripped.
    """
    text = re.sub(r"\[([^\]]*)\]\(([^)]+)\)", r"\1", text)   # links -> label
    out = []
    for part in TOKEN.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            out.append((part[2:-2], True, False, None))
        elif part.startswith("<sup>") and part.endswith("</sup>"):
            out.append((part[5:-6], False, False, "super"))
        elif part.startswith("<sub>") and part.endswith("</sub>"):
            out.append((part[5:-6], False, False, "sub"))
        elif part.startswith("*") and part.endswith("*"):
            out.append((part[1:-1], False, True, None))
        elif part.startswith("`") and part.endswith("`"):
            out.append((part[1:-1], False, False, None))
        else:
            out.append((part, False, False, None))
    return out


HEADNUM = re.compile(r"^(\d+(?:\.\d+)*\.?)(\s+)(.*)$")


def heading_parts(lvl, txt):
    """Split a heading into (text, bold, italic) runs per the JHSS scheme.

    The Quick Guide and the General Format section name the same ladder one
    level apart; taken together they give:
        ##   section        number bold,  title bold
        ###  subsection     number bold,  title italic
        #### sub-subsection number plain, title italic
    """
    num_bold = lvl <= 3
    title_bold = lvl <= 2
    title_ital = lvl >= 3
    m = HEADNUM.match(txt)
    if not m:
        return [(txt, title_bold, title_ital)]
    num, gap, title = m.groups()
    return [(num + gap, num_bold, False), (title, title_bold, title_ital)]


def _table_rules(tbl, header_rows=1):
    """JHSS: hide all table lines except the rule under the column headings."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tblPr = tbl._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        borders.append(el)
    tblPr.append(borders)
    for row in tbl.rows[:header_rows]:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcB = OxmlElement("w:tcBorders")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "8")
            tcB.append(bottom)
            tcPr.append(tcB)


def _script(run, sc):
    """Apply super/subscript to a python-docx run."""
    if sc == "super":
        run.font.superscript = True
    elif sc == "sub":
        run.font.subscript = True


def figpath(rel):
    p = os.path.join(FIGROOT, rel)
    return p if os.path.exists(p) else None


# ---------------------------------------------------------------- docx

def build_docx(blocks, dest):
    import docx
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    doc = docx.Document()

    # JHSS quick guide: one column, Times New Roman 12, line spacing 1.15,
    # 1 in left/right margins, 0.5 in top/bottom, paragraphs separated by a
    # blank line and never indented.
    for sec in doc.sections:
        sec.left_margin = sec.right_margin = Inches(1)
        sec.top_margin = sec.bottom_margin = Inches(0.5)

    for st, sz in (("Normal", 12), ("Title", 16), ("Heading 1", 14),
                   ("Heading 2", 13), ("Heading 3", 12), ("Heading 4", 12)):
        try:
            s = doc.styles[st]
            s.font.name = "Times New Roman"
            s.font.size = Pt(sz)
            if st.startswith("Heading") or st == "Title":
                s.font.color.rgb = RGBColor(0, 0, 0)
            # Word's default Title style carries a blue 1pt rule underneath it;
            # drop any style-level border so no decorative bar is drawn.
            pPr = s.element.get_or_add_pPr()
            for bdr in pPr.findall(qn("w:pBdr")):
                pPr.remove(bdr)
        except KeyError:
            pass

    pf = doc.styles["Normal"].paragraph_format
    pf.line_spacing = 1.15
    pf.space_after = Pt(10)
    pf.first_line_indent = Inches(0)

    for kind, payload in blocks:
        if kind == "h":
            lvl, txt = payload
            p = doc.add_paragraph()
            p.style = doc.styles["Title" if lvl == 1 else f"Heading {min(lvl - 1, 4)}"]
            for t, b, it in heading_parts(lvl, txt):
                r = p.add_run(t)
                r.bold, r.italic = b, it
        elif kind == "p":
            p = doc.add_paragraph()
            # JHSS: figure titles left-justified at size 10 below the figure;
            # table titles left-justified at size 12 above the table.
            is_fig = re.match(r"\*\*Figure \d+\.\*\*", payload)
            is_tab = re.match(r"\*\*Table \d+\.\*\*", payload)
            p.alignment = (WD_ALIGN_PARAGRAPH.LEFT if (is_fig or is_tab)
                           else WD_ALIGN_PARAGRAPH.JUSTIFY)
            for t, b, it, sc in runs(payload):
                r = p.add_run(t)
                r.bold, r.italic = b, it
                _script(r, sc)
                if is_fig:
                    r.font.size = Pt(10)
        elif kind == "li":
            marker, txt = payload
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.add_run(marker + " ")
            for t, b, it, sc in runs(txt):
                r = p.add_run(t)
                r.bold, r.italic = b, it
                _script(r, sc)
        elif kind == "quote":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.35)
            p.paragraph_format.right_indent = Inches(0.2)
            for t_, b_, i_, sc in runs(payload):
                r = p.add_run(t_)
                r.bold, r.italic = b_, True
                _script(r, sc)
        elif kind == "img":
            fp = figpath(payload)
            if fp:
                doc.add_picture(fp, width=Inches(6.2))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif kind == "table":
            rows = payload
            ncol = max(len(r) for r in rows)
            tbl = doc.add_table(rows=0, cols=ncol)
            tbl.style = "Table Grid"
            _rule_after = True
            for ri, row in enumerate(rows):
                cells = tbl.add_row().cells
                for ci in range(ncol):
                    cell = cells[ci]
                    cell.text = ""
                    para = cell.paragraphs[0]
                    for t, b, it, sc in runs(row[ci] if ci < len(row) else ""):
                        r = para.add_run(t)
                        r.bold = b or ri == 0
                        r.italic = it
                        _script(r, sc)
                        r.font.size = Pt(9)
            _table_rules(tbl)
            doc.add_paragraph()
    doc.save(dest)
    return dest


# ---------------------------------------------------------------- pdf

def build_pdf(blocks, dest):
    from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Image,
                                    Table, TableStyle, KeepTogether)
    from PIL import Image as PILImage

    pdfmetrics.registerFont(TTFont("DJ", f"{DEJAVU}/DejaVuSerif.ttf"))
    pdfmetrics.registerFont(TTFont("DJ-B", f"{DEJAVU}/DejaVuSerif-Bold.ttf"))
    pdfmetrics.registerFont(TTFont("DJS", f"{DEJAVU}/DejaVuSans.ttf"))
    pdfmetrics.registerFont(TTFont("DJS-B", f"{DEJAVU}/DejaVuSans-Bold.ttf"))
    pdfmetrics.registerFontFamily("DJ", normal="DJ", bold="DJ-B", italic="DJ", boldItalic="DJ-B")

    body = ParagraphStyle("body", fontName="DJ", fontSize=9.4, leading=13.4,
                          alignment=TA_JUSTIFY, spaceAfter=6)
    title = ParagraphStyle("title", fontName="DJS-B", fontSize=16, leading=20,
                           alignment=TA_CENTER, spaceAfter=14)
    heads = {2: ParagraphStyle("h2", fontName="DJS-B", fontSize=12.5, leading=16,
                               spaceBefore=13, spaceAfter=6),
             3: ParagraphStyle("h3", fontName="DJS-B", fontSize=10.8, leading=14,
                               spaceBefore=10, spaceAfter=5),
             4: ParagraphStyle("h4", fontName="DJS-B", fontSize=9.8, leading=13,
                               spaceBefore=8, spaceAfter=4)}
    cellst = ParagraphStyle("cell", fontName="DJ", fontSize=6.9, leading=8.6)
    cellhd = ParagraphStyle("cellh", fontName="DJS-B", fontSize=6.9, leading=8.6)
    listst = ParagraphStyle("li", parent=body, leftIndent=16)
    quotest = ParagraphStyle("quote", parent=body, leftIndent=22, rightIndent=14,
                             textColor=colors.HexColor("#333333"), spaceBefore=4, spaceAfter=8)

    def esc(text):
        out = ""
        for t, b, it, sc in runs(text):
            t = html.escape(t)
            if b:
                t = f"<font name='DJ-B'>{t}</font>"
            if it:
                t = f"<i>{t}</i>"
            if sc == "super":
                t = f"<super>{t}</super>"
            elif sc == "sub":
                t = f"<sub>{t}</sub>"
            out += t
        return out

    story, avail = [], LETTER[0] - 1.5 * inch
    for kind, payload in blocks:
        if kind == "h":
            lvl, txt = payload
            frag = ""
            for t, b, it in heading_parts(lvl, txt):
                t = html.escape(t)
                if b:
                    t = f"<font name='DJS-B'>{t}</font>"
                if it:
                    t = f"<i>{t}</i>"
                frag += t
            story.append(Paragraph(frag, title if lvl == 1 else heads.get(min(lvl, 4), heads[4])))
        elif kind == "p":
            story.append(Paragraph(esc(payload), body))
        elif kind == "li":
            marker, txt = payload
            story.append(Paragraph(f"{html.escape(marker)} {esc(txt)}", listst))
        elif kind == "quote":
            story.append(Paragraph(f"<i>{esc(payload)}</i>", quotest))
        elif kind == "img":
            fp = figpath(payload)
            if not fp:
                continue
            with PILImage.open(fp) as im:
                w, h = im.size
            scale = min(avail / w, (LETTER[1] - 2.6 * inch) / h)
            story.append(Spacer(1, 5))
            story.append(Image(fp, width=w * scale, height=h * scale))
            story.append(Spacer(1, 5))
        elif kind == "table":
            rows = payload
            ncol = max(len(r) for r in rows)
            data = [[Paragraph(esc(r[c] if c < len(r) else ""), cellhd if i == 0 else cellst)
                     for c in range(ncol)] for i, r in enumerate(rows)]
            t = Table(data, colWidths=[avail / ncol] * ncol, repeatRows=1)
            t.setStyle(TableStyle([
                ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#999999")),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#ededed")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 3),
                ("RIGHTPADDING", (0, 0), (-1, -1), 3),
                ("TOPPADDING", (0, 0), (-1, -1), 2.5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 2.5)]))
            story.append(Spacer(1, 4))
            story.append(t)
            story.append(Spacer(1, 8))

    SimpleDocTemplate(dest, pagesize=LETTER,
                      leftMargin=0.75 * inch, rightMargin=0.75 * inch,
                      topMargin=0.8 * inch, bottomMargin=0.8 * inch,
                      title="Revised manuscript").build(story)
    return dest


FIGCAP = re.compile(r"^\*\*Figure\s+\d+\.\*\*")


def split_figures(blocks):
    """Return (blocks without figures/captions, caption blocks).

    Only the caption half is used now: the manuscript keeps its figures, and the
    legends are emitted separately as JHSS requires."""
    body, legends = [], []
    for kind, payload in blocks:
        if kind == "img":
            continue
        if kind == "p" and FIGCAP.match(payload):
            legends.append((kind, payload))
            continue
        body.append((kind, payload))
    return body, legends


if __name__ == "__main__":
    if len(sys.argv) > 2 and sys.argv[2].endswith(".md"):
        SRC = sys.argv[2]
    blocks = parse(open(SRC, encoding="utf-8").read())
    kinds = {}
    for k, _ in blocks:
        kinds[k] = kinds.get(k, 0) + 1
    print("parsed blocks:", kinds)
    which = sys.argv[1] if len(sys.argv) > 1 else "both"
    stem = os.path.splitext(os.path.basename(SRC))[0]

    if which in ("docx", "both"):
        print("docx ->", build_docx(blocks, f"{OUTDIR}/{stem}.docx"))
    if which in ("pdf", "both"):
        print("pdf  ->", build_pdf(blocks, f"{OUTDIR}/{stem}.pdf"))

    # JHSS requires figure legends as their own file, alongside the manuscript
    # (which keeps its figures) and the separate JPEGs. Emitted whenever the
    # source actually contains figure captions.
    _, legends = split_figures(blocks)
    if legends:
        head = [("h", (1, "Figure legends"))]
        print("legends ->", build_docx(head + legends, f"{OUTDIR}/figure_legends.docx"))
